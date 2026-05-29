from docx import Document
from docx.shared import Pt

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

doc.add_heading('Felicia 360: Documentacao Tecnica', level=0)
doc.add_paragraph('Pricing Operacoes | Versao 3.0 | 21/05/2026')
doc.add_paragraph('')

doc.add_heading('1. Visao Geral', level=1)
doc.add_paragraph(
    'Dashboard em Google Apps Script + React que consolida a visao economica do cliente Stone, '
    'unificando credito, adquirencia e banking numa interface unica para tomada de decisao.'
)
doc.add_paragraph('URL: https://script.google.com/a/macros/stone.com.br/s/AKfycbxTNqpYwEBxdXLg-qmcNGjM-agAFAeLAt4YEgZOvh7m7KnRaLwuMtuGmgg__h4TrrmI/exec')
doc.add_paragraph('Projeto GAS: 1qKErmyBogImZKfioaCkDPcowFYkTGQnFDQbpZk1kSKe_zXk_N0vqthGU')

doc.add_heading('2. Stack Tecnica', level=1)
for line in [
    'Frontend: React 18 + Tailwind CSS + Recharts',
    'Build: Vite + vite-plugin-singlefile (HTML unico inline)',
    'Backend: Google Apps Script com BigQuery Advanced Service',
    'Deploy: clasp push + clasp deploy -i <ID>',
    'Acesso: DOMAIN (qualquer @stone.com.br)',
    'Execucao: USER_DEPLOYING (queries rodam com credenciais do deployer)',
]:
    doc.add_paragraph(line)

doc.add_heading('3. Componentes do Dashboard', level=1)
components = [
    ('SearchBar', 'Input de CNPJ/CPF com limpeza de formatacao. Dispara 7 queries paralelas ao buscar.'),
    ('Informacoes do Cliente', 'Grid 2 colunas: Adquirencia (documento, afiliacoes, MCC, SS1/SS3/SS5, produtos) + Status da Conta (KYC, rating, tipo conta).'),
    ('Credito: Lifetime (VP)', 'NII, Risk Adj NII, NPV a valor presente. Resumo de desembolsos (qtd, total, adimplencia). Tooltip com definicoes.'),
    ('Adquirencia: Media 3m', 'TPV, Margem, Margem/TPV, Net MDR%, TKR nCOF. Media dos 3 ultimos meses fechados (exclui mes corrente). Aviso quando TPV < R$1.000.'),
    ('Banco: Media 3m', 'Saldo Conta, Saldo Reservas, Boletos Emitidos, Boletos Liquidados, Volume Boleto (R$). Tooltip com definicoes.'),
    ('Adquirencia: Detalhado Mensal', 'Tabela com 19 colunas. Dropdown de filtro de colunas. Mes formatado MM/YYYY. Valores negativos em vermelho.'),
    ('Ofertas de Credito', 'Cartao (limites, atraso) + Demais Ofertas (filtro de status, Cancelled/Denied ocultos por default) + Desembolso (contratos ordenados por data DESC).'),
    ('Fluxo de Caixa Mensal', 'Grafico com 5 linhas credito (azul) + adquirencia (verde). Zoom via scroll + Brush. Filtros toggle. Linha "Hoje" no mes atual.'),
    ('Monitor Admin', 'Visivel apenas para ayran.maduro@stone.com.br. Heartbeat 30s. Painel com usuarios online/offline e historico.'),
]
for name, desc in components:
    p = doc.add_paragraph()
    run = p.add_run(name + ': ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('4. Queries BigQuery', level=1)
doc.add_paragraph('Todas executadas via BigQuery Advanced Service. Projeto BQ: sbj7ujlwjbsknn8v396xaahlf4ogck.')

# Query 1
doc.add_heading('4.1 getStatusCredito(doc)', level=2)
doc.add_paragraph('Tabela: Dias_Credit.base_felicia')
doc.add_paragraph('Descricao: Status KYC, rating, ofertas, cartao, desembolsos')
doc.add_paragraph('Grao: 1 linha por oferta/contrato | Refresh: Diario')
p = doc.add_paragraph()
r = p.add_run('SELECT * FROM base_felicia WHERE CNPJ = @doc')
r.font.name = 'Consolas'
r.font.size = Pt(8)

# Query 2
doc.add_heading('4.2 getNpvCredito(doc)', level=2)
doc.add_paragraph('Tabelas: credit_pricing.npv_kgiro + credit_policy_studies.credit_portfolio')
doc.add_paragraph('Descricao: NII, Risk Adj NII, NPV (lifetime, valor presente)')
doc.add_paragraph('Grao: 1 linha por documento | Refresh: Periodico (run_at)')
sql2 = """WITH loans AS (
  SELECT DISTINCT LoanId, documento
  FROM credit_portfolio WHERE documento = @doc
)
SELECT
  ROUND(SUM(financial_income_net * discount_factor)
    - SUM(funding_cost * discount_factor)
    - SUM(capital_cost * discount_factor), 2) AS nii,
  ROUND(SUM(financial_income_net * discount_factor)
    - SUM(funding_cost * discount_factor)
    - SUM(capital_cost * discount_factor)
    - SUM(pdd_result * discount_factor), 2) AS risk_adj_nii,
  ROUND(SUM(pv_cf), 2) AS npv
FROM npv_kgiro a JOIN loans b ON a.loanid = b.LoanId
WHERE run_at = (SELECT MAX(run_at) FROM npv_kgiro)
GROUP BY b.documento"""
p = doc.add_paragraph()
r = p.add_run(sql2)
r.font.name = 'Consolas'
r.font.size = Pt(8)

# Query 3
doc.add_heading('4.3 getPnlAdquirencia(doc)', level=2)
doc.add_paragraph('Tabela: Dias_PnL.PnL_Dashs_part')
doc.add_paragraph('Descricao: Serie mensal (12m): TPV, NetMDR, Floating, Aluguel, RAV, TED, Pix, Gateway, Rcta NetCOF, COGs, Margem + %')
doc.add_paragraph('Grao: documento x mes | Refresh: Diario')
sql3 = """SELECT FORMAT_DATE('%Y-%m', Dt_Month) AS mes,
  ROUND(SUM(GMV), 2) AS tpv,
  ROUND(SUM(Net_MDR_Stone), 2) AS net_mdr,
  ROUND(SUM(Floating_Stn), 2) AS floating_conta,
  ROUND(SUM(Rcta_Aluguel), 2) AS aluguel,
  ROUND(SUM(Margem_RAV_STN), 2) AS net_rav,
  ROUND(SUM(Receita_TED), 2) AS rcta_ted,
  ROUND(SUM(Receita_Pix_Geral), 2) AS pix_rcta,
  ROUND(SUM(Rcta_gateway), 2) AS gateway,
  ROUND(SUM(Receita_Net_COF), 2) AS receita_net_cof,
  ROUND(SUM(custo_servir_Total) * (-1), 2) AS cogs,
  ROUND(SUM(Margem_Query), 2) AS margem
  -- + percentuais via SAFE_DIVIDE
FROM PnL_Dashs_part
WHERE ClientCNPJorCPF = @doc
  AND Dt_Month >= DATE_SUB(CURRENT_DATE(), INTERVAL 12 MONTH)
GROUP BY Dt_Month ORDER BY Dt_Month DESC"""
p = doc.add_paragraph()
r = p.add_run(sql3)
r.font.name = 'Consolas'
r.font.size = Pt(8)

# Query 4
doc.add_heading('4.4 getFluxoCreditoMensal(doc)', level=2)
doc.add_paragraph('Tabelas: credit_pricing.npv_kgiro + credit_portfolio')
doc.add_paragraph('Descricao: Serie mensal de credito: receita juros, funding, capital, PDD, NII, Risk Adj NII')
doc.add_paragraph('Grao: documento x mes (pode ir ate 2031) | Refresh: Periodico')
sql4 = """WITH loans AS (
  SELECT DISTINCT LoanId, documento
  FROM credit_portfolio WHERE documento = @doc
)
SELECT FORMAT_DATE('%Y-%m', reference_date) AS mes,
  ROUND(SUM(financial_income_net), 2) AS receita_juros,
  ROUND(SUM(funding_cost), 2) AS funding_cost,
  ROUND(SUM(capital_cost), 2) AS capital_cost,
  ROUND(SUM(pdd_result), 2) AS pdd_result,
  ROUND(SUM(financial_income_net - funding_cost - capital_cost), 2) AS nii,
  ROUND(SUM(financial_income_net - funding_cost - capital_cost - pdd_result), 2) AS risk_adj_nii
FROM npv_kgiro a JOIN loans b ON a.loanid = b.LoanId
WHERE run_at = (SELECT MAX(run_at) FROM npv_kgiro)
GROUP BY reference_date ORDER BY reference_date"""
p = doc.add_paragraph()
r = p.add_run(sql4)
r.font.name = 'Consolas'
r.font.size = Pt(8)

# Query 5
doc.add_heading('4.5 getInfoClienteAdq(doc)', level=2)
doc.add_paragraph('Tabela: Dias_PnL.PnL_Dashs_part')
doc.add_paragraph('Descricao: Dados cadastrais: afiliacoes, nome, MCC, canal, regional, polo, produtos')
doc.add_paragraph('Grao: 1 linha agregada | Refresh: Diario')
sql5 = """SELECT
  STRING_AGG(DISTINCT CAST(SC AS STRING), ', ') AS stonecodes,
  MAX(ClientName) AS client_name,
  MAX(CAST(MCC AS STRING)) AS mcc,
  MAX(SalesStructureNameLevel1) AS nivel1,
  MAX(SalesStructureNameLevel3) AS nivel3,
  MAX(SalesStructureNameLevel5) AS nivel5,
  STRING_AGG(DISTINCT CompanyName, ', ') AS produtos
FROM PnL_Dashs_part
WHERE ClientCNPJorCPF = @doc
  AND Dt_Month >= DATE_SUB(CURRENT_DATE(), INTERVAL 3 MONTH)"""
p = doc.add_paragraph()
r = p.add_run(sql5)
r.font.name = 'Consolas'
r.font.size = Pt(8)

# Query 6
doc.add_heading('4.6 getBancoMedia(doc)', level=2)
doc.add_paragraph('Tabela: Dias_PnL.resumo_conta_3M')
doc.add_paragraph('Descricao: Media 3 meses: saldo conta, saldo reservas, boletos emitidos/liquidados, volume boleto')
doc.add_paragraph('Grao: 1 linha agregada | Refresh: Diario')
sql6 = """WITH ft AS (
  SELECT mes,
    SUM(Media_Saldo_Conta_Visao_Cliente) AS saldo_conta,
    SUM(Media_Saldo_Reservas) AS saldo_reservas,
    SUM(IFNULL(Quantidade_total_emitida, 0)) AS boleto_emitido,
    SUM(IFNULL(Quantidade_de_boletos_liquidados, 0)) AS boleto_liquidado,
    SUM(IFNULL(Valor_total_liquidado_pago, 0)) AS volume_boleto_liquidado
  FROM resumo_conta_3M WHERE Documento = @doc
  GROUP BY mes
)
SELECT
  ROUND(AVG(saldo_conta), 2) AS saldo_conta,
  ROUND(AVG(saldo_reservas), 2) AS saldo_reservas,
  ROUND(AVG(boleto_emitido), 1) AS boleto_emitido,
  ROUND(AVG(boleto_liquidado), 1) AS boleto_liquidado,
  ROUND(AVG(volume_boleto_liquidado), 2) AS volume_boleto
FROM ft"""
p = doc.add_paragraph()
r = p.add_run(sql6)
r.font.name = 'Consolas'
r.font.size = Pt(8)

# Metricas
doc.add_heading('5. Definicoes de Metricas', level=1)
metricas = [
    ('NII', 'Receita de Juros - Custo de Funding, trazendo todo fluxo a valor presente'),
    ('Risk Adj NII', 'NII - Custo de Risco (PDD), fluxo a VP'),
    ('NPV', 'Risk Adj NII - Custo Operacional - CAC'),
    ('TPV', 'Volume Total de Pagamentos (GMV)'),
    ('Net MDR %', '[MDR - (IC + Fee)] / cTPV'),
    ('TKR nCOF', 'Receita nCOF / TPV'),
    ('Margem', 'Receita nCOF - COGs'),
    ('Media 3m', 'Media dos 3 ultimos meses fechados (exclui mes corrente)'),
]
for name, desc in metricas:
    p = doc.add_paragraph()
    run = p.add_run(name + ': ')
    run.bold = True
    p.add_run(desc)

# Gotchas
doc.add_heading('6. Gotchas e Decisoes Tecnicas', level=1)
gotchas = [
    'Campo Gateway na PnL_Dashs_part e Rcta_gateway (nao Rcta_gateway_sum)',
    'Datas do BQ Advanced Service vem como epoch seconds em notacao cientifica (ex: 1.729E9). Frontend converte via parseFloat * 1000',
    'Media 3m exclui mes corrente (aberto/incompleto) e ordena DESC antes de pegar os 3 primeiros',
    'Adimplencia: "Em dia" e "Encerrado" sao adimplentes. So entra como atraso faixas tipo "15 a 30" etc.',
    'Build em /c/temp/felicia-build/ (path sem espacos). Google Drive com espacos quebra o Vite',
    'Deploy sempre no mesmo ID para manter URL fixa',
    'Monitor admin visivel apenas para ayran.maduro@stone.com.br',
]
for g in gotchas:
    doc.add_paragraph(g, style='List Bullet')

doc.save('documentos/felicia-360-documentacao-tecnica.docx')
print('Documento salvo com sucesso')
